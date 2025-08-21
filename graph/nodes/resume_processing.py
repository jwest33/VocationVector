"""
Resume Processing Node - Extracts structured information from resumes for job matching

This module uses llama.cpp with Qwen3-4B model for LLM processing.
Ensure llama.cpp server is running at http://localhost:8000/v1

To start the server:
./server -m models/qwen3-4b-instruct-2507-f16.gguf --host 0.0.0.0 --port 8000

Environment variables (optional):
- LLM_MODEL: Model name (default: qwen3-4b-instruct-2507-f16)
- OPENAI_BASE_URL: Server URL (default: http://localhost:8000/v1)
"""

from __future__ import annotations
import os
import json
import logging
import re
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

# Set up module logger
logger = logging.getLogger(__name__)

# LLM client (OpenAI-compatible)
from openai import OpenAI
from tenacity import retry, stop_after_attempt, wait_exponential

# Universal template
from graph.vocation_template import (
    VocationTemplate,
    SkillRequirement,
    ExperienceRequirement,
    EducationRequirement,
    LocationRequirement,
    CompensationRequirement,
    CultureFit,
    WorkArrangement,
    EmploymentType,
    JobSeekerPreferences
)

# Validation
from pydantic import BaseModel, Field, field_validator

# PDF handling (optional)
try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logging.info("PyPDF2 not available, PDF support disabled")

# DOCX handling (optional)
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.info("python-docx not available, DOCX support disabled")


# -------------------------
# Configuration
# -------------------------

class ResumeConfig:
    """Configuration for resume processing"""
    
    # LLM settings for llama.cpp
    DEFAULT_LLM_MODEL: str = "qwen3-4b-instruct-2507-f16"  # For llama.cpp, this is ignored but required by OpenAI client
    LLM_TEMPERATURE: float = 0.1
    LLM_MAX_TOKENS: int = 262144  # Increased for more comprehensive responses with Qwen3-4B
    LLM_MAX_RETRIES: int = 3
    LLM_TIMEOUT: int = 600  # 10 minutes timeout for large prompts and comprehensive extraction
    
    # llama.cpp server defaults
    DEFAULT_BASE_URL: str = "http://localhost:8000/v1"  # Standard llama.cpp port
    
    # Text processing
    MAX_RESUME_CHARS: int = 15000  # Limit for LLM processing
    
    # Resume sections to extract (maps to job posting fields)
    RESUME_SECTIONS: List[str] = [
        "contact_info",
        "professional_summary",
        "technical_skills",
        "programming_languages",
        "frameworks_tools",
        "work_experience",
        "education",
        "certifications",
        "achievements",
        "salary_expectations",
        "location_preferences",
        "work_authorization",
        "availability"
    ]
    
    # Map resume fields to job posting fields
    FIELD_MAPPING: Dict[str, List[str]] = {
        "technical_skills": ["requirements", "qualifications", "must_have"],
        "work_experience": ["experience", "responsibilities", "duties"],
        "education": ["education", "degree", "qualifications"],
        "certifications": ["certifications", "licenses", "credentials"],
        "location_preferences": ["location", "remote_policy", "work_arrangement"],
        "salary_expectations": ["compensation", "salary", "pay"],
        "availability": ["start_date", "availability", "notice_period"]
    }
    
    @classmethod
    def from_dict(cls, config_dict: Dict[str, Any]) -> 'ResumeConfig':
        """Create config from dictionary"""
        config = cls()
        for key, value in config_dict.items():
            if hasattr(config, key):
                setattr(config, key, value)
        return config


# -------------------------
# Data Models
# -------------------------

@dataclass
class ContactInfo:
    """Contact information from resume"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None


@dataclass
class WorkExperience:
    """Single work experience entry"""
    title: str
    company: str
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    duration: Optional[str] = None
    responsibilities: List[str] = field(default_factory=list)
    achievements: List[str] = field(default_factory=list)
    technologies: List[str] = field(default_factory=list)


@dataclass
class Education:
    """Education entry"""
    degree: str
    institution: str
    location: Optional[str] = None
    graduation_date: Optional[str] = None
    gpa: Optional[str] = None
    relevant_courses: List[str] = field(default_factory=list)


@dataclass
class Certification:
    """Professional certification"""
    name: str
    issuer: Optional[str] = None
    date_obtained: Optional[str] = None
    expiry_date: Optional[str] = None
    credential_id: Optional[str] = None


@dataclass
class ProcessedResume:
    """Complete processed resume structure"""
    contact_info: ContactInfo
    professional_summary: str
    technical_skills: List[str]
    programming_languages: List[str]
    frameworks_tools: List[str]
    work_experience: List[WorkExperience]
    education: List[Education]
    certifications: List[Certification]
    achievements: List[str]
    years_of_experience: Optional[float] = None
    salary_expectations: Optional[Dict[str, Any]] = None
    location_preferences: Optional[Dict[str, Any]] = None
    work_authorization: Optional[str] = None
    availability: Optional[str] = None
    raw_text: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class ResumeExtraction(BaseModel):
    """Validated LLM response for resume extraction"""
    contact_info: Dict[str, Optional[str]] = Field(default_factory=dict)
    professional_summary: str = Field(default="", max_length=1000)
    technical_skills: List[str] = Field(default_factory=list)
    programming_languages: List[str] = Field(default_factory=list)
    frameworks_tools: List[str] = Field(default_factory=list)
    work_experience: List[Dict[str, Any]] = Field(default_factory=list)
    education: List[Dict[str, Any]] = Field(default_factory=list)
    certifications: List[Dict[str, Any]] = Field(default_factory=list)
    achievements: List[str] = Field(default_factory=list)
    years_of_experience: Optional[float] = None
    salary_expectations: Optional[Dict[str, Any]] = None
    location_preferences: Optional[Dict[str, Any]] = None
    work_authorization: Optional[str] = None
    availability: Optional[str] = None
    
    @field_validator('years_of_experience')
    def validate_experience(cls, v):
        if v is not None and v < 0:
            return 0
        return v


# -------------------------
# Text Extraction
# -------------------------

class ResumeTextExtractor:
    """Extract text from various resume formats"""
    
    @staticmethod
    def extract_from_file(file_path: str) -> str:
        """Extract text from a resume file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.txt':
            return ResumeTextExtractor._extract_from_txt(file_path)
        elif extension == '.pdf' and HAS_PYPDF:
            return ResumeTextExtractor._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc'] and HAS_DOCX:
            return ResumeTextExtractor._extract_from_docx(file_path)
        else:
            # Try to read as text file
            try:
                return ResumeTextExtractor._extract_from_txt(file_path)
            except Exception as e:
                raise ValueError(f"Unsupported file format or missing dependencies: {extension}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from a text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file"""
        if not HAS_PYPDF:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        
        return '\n'.join(text)
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file"""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX processing")
        
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return '\n'.join(text)


# -------------------------
# LLM Processing
# -------------------------

class ResumeLLMProcessor:
    """Process resume text using LLM to extract structured information"""
    
    def __init__(
        self,
        model: str = None,
        base_url: str = None,
        api_key: str = None,
        config: ResumeConfig = None
    ):
        self.config = config or ResumeConfig()
        self.model = model or os.getenv("LLM_MODEL", self.config.DEFAULT_LLM_MODEL)
        self.base_url = base_url or os.getenv("OPENAI_BASE_URL", self.config.DEFAULT_BASE_URL)
        # llama.cpp doesn't need API key, but OpenAI client requires something
        self.api_key = api_key or os.getenv("OPENAI_API_KEY", "dummy-key")
        
        self.client = OpenAI(
            base_url=self.base_url,
            api_key=self.api_key
        )
        
        logging.info(f"Initialized Resume LLM processor with model: {self.model} at {self.base_url}")
    
    def _analyze_resume_structure(self, resume_text: str) -> Dict[str, Any]:
        """First pass: Analyze resume to understand its structure"""
        
        # Take a sample of the resume for analysis
        sample = resume_text[:1500] if len(resume_text) > 1500 else resume_text
        
        system_prompt = "Identify resume sections. Return JSON only."
        
        # Simpler analysis prompt
        user_prompt = f"""Analyze this resume:

{sample}

Return JSON:
{{
  "has_summary": true,
  "has_skills": true,
  "has_experience": true,
  "has_education": true,
  "detected_field": "payroll"
}}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            analysis = self._parse_llm_response(response)
            logging.info(f"Resume analysis: {analysis}")
            return analysis
        except Exception as e:
            logging.warning(f"Resume analysis failed: {e}, using defaults")
            return {
                "has_contact_info": True,
                "has_summary": True,
                "has_skills": True,
                "has_experience": True,
                "detected_field": "general"
            }
    
    def _build_dynamic_extraction_prompt(self, resume_text: str, analysis: Dict[str, Any]) -> Tuple[str, str]:
        """Build extraction prompt based on resume analysis"""
        
        system_prompt = "Extract resume data. Return complete JSON."
        
        # Truncate more aggressively for better processing
        if len(resume_text) > 3000:
            resume_text = resume_text[:3000]
        
        # Detect field type from text
        field = "payroll" if "payroll" in resume_text.lower() else "general"
        
        # Simpler, more focused prompt
        user_prompt = f"""Extract from resume:

{resume_text}

Return this JSON structure exactly:
{{
  "contact_info": {{
    "name": "full name",
    "email": "email",
    "phone": "phone",
    "location": "city, state"
  }},
  "professional_summary": "first paragraph",
  "technical_skills": ["skill1", "skill2"],
  "programming_languages": [],
  "frameworks_tools": ["ADP", "Excel"],
  "work_experience": [
    {{
      "title": "job title",
      "company": "company",
      "start_date": "start",
      "end_date": "end",
      "responsibilities": ["duty1"],
      "technologies": []
    }}
  ],
  "education": [
    {{
      "degree": "degree",
      "institution": "school",
      "graduation_date": "year"
    }}
  ],
  "certifications": [],
  "years_of_experience": 7,
  "work_authorization": "",
  "availability": ""
}}"""
        
        return system_prompt, user_prompt
    
    def _build_extraction_prompt(self, resume_text: str) -> Tuple[str, str]:
        """Two-pass approach: analyze then extract"""
        
        # First pass: analyze structure
        analysis = self._analyze_resume_structure(resume_text)
        
        # Second pass: build dynamic prompt
        return self._build_dynamic_extraction_prompt(resume_text, analysis)
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def _call_llm(self, system_prompt: str, user_prompt: str) -> str:
        """Call LLM with retry logic"""
        import time
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                # Simplified call for better compatibility
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=self.config.LLM_TEMPERATURE,
                    max_tokens=10000,  # Reduced from 262144 for better compatibility
                    timeout=300  # 5 minutes timeout
                )
                
                # Check if we got a valid response
                if not response.choices or not response.choices[0].message.content:
                    logging.warning(f"Empty response on attempt {attempt + 1}, retrying...")
                    if attempt < max_retries - 1:
                        time.sleep(3)  # Wait for server to warm up
                        continue
                    else:
                        # On last attempt, return empty JSON
                        logging.error("All attempts failed, returning empty result")
                        return "{}"
                
                content = response.choices[0].message.content
                logging.debug(f"LLM response: {content[:200] if content else 'Empty'}...")
                return content
                
            except Exception as e:
                logging.error(f"LLM call failed on attempt {attempt + 1}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(3)
                    continue
                else:
                    # On last attempt, return empty JSON
                    return "{}"
        
        return "{}"
    
    def _extract_contact_info(self, resume_text: str) -> Dict[str, Any]:
        """Extract contact information from resume"""
        system_prompt = "You are a resume parser. Extract ONLY explicitly stated contact information. Do NOT infer or assume information."
        
        # Use full resume text for contact extraction
        resume_snippet = resume_text
        
        user_prompt = f"""Extract ONLY the explicitly stated contact information from this resume. 
IMPORTANT RULES:
- Only extract information that is clearly the person's contact details
- Do NOT use job locations as the person's location
- Do NOT make assumptions or inferences
- If information is not explicitly stated as contact info, use null

Return ONLY valid JSON with these fields (use null for missing information):
{{
    "name": "full name if stated",
    "email": "email address if stated", 
    "phone": "phone number if stated",
    "location": "ONLY if explicitly stated as person's location, NOT job location",
    "linkedin": "LinkedIn URL if stated",
    "github": "GitHub URL if stated"
}}

Resume text:
{resume_snippet}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            result = self._parse_llm_response(response)
            # Ensure we have the expected structure
            if isinstance(result, dict):
                return result
            return {"name": None, "email": None, "phone": None, "location": None}
        except Exception as e:
            logging.error(f"Failed to extract contact info: {e}")
            return {"name": None, "email": None, "phone": None, "location": None}
    
    def _extract_summary(self, resume_text: str) -> str:
        """Extract professional summary from resume"""
        system_prompt = "You are a resume parser. Extract the professional summary."
        
        # Use full resume text for summary extraction
        resume_snippet = resume_text
        
        user_prompt = f"""Extract the professional summary or objective from this resume. Return ONLY valid JSON:
{{
    "summary": "the professional summary text",
    "has_summary": true/false
}}

Resume text:
{resume_snippet}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            result = self._parse_llm_response(response)
            if isinstance(result, dict) and "summary" in result:
                return result["summary"][:1000]  # Limit length
            # If it's a plain string response
            if isinstance(result, str):
                return result[:1000]
            return ""
        except Exception as e:
            logging.error(f"Failed to extract summary: {e}")
            return ""
    
    def _extract_work_experience(self, resume_text: str) -> List[Dict[str, Any]]:
        """Extract work experience from resume"""
        system_prompt = "You are a resume parser. Extract work experience."
        
        # Use full resume text for work experience extraction
        resume_snippet = resume_text
        
        user_prompt = f"""Extract the work experience from this resume. Return ONLY valid JSON:
{{
    "experience": [
        {{
            "company": "company name",
            "role": "job title",
            "duration": "time period",
            "location": "location if mentioned",
            "description": ["bullet point 1", "bullet point 2"]
        }}
    ]
}}

Resume text:
{resume_snippet}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            result = self._parse_llm_response(response)
            # Handle if response is dict with 'experience' key (matching our prompt)
            if isinstance(result, dict) and 'experience' in result:
                # Convert to expected format
                experience = []
                for exp in result['experience']:
                    experience.append({
                        "title": exp.get("role", ""),
                        "company": exp.get("company", ""),
                        "start_date": exp.get("duration", "").split("-")[0].strip() if "-" in exp.get("duration", "") else "",
                        "end_date": exp.get("duration", "").split("-")[1].strip() if "-" in exp.get("duration", "") else exp.get("duration", ""),
                        "responsibilities": exp.get("description", []),
                        "technologies": []
                    })
                return experience
            elif isinstance(result, list):
                return result
            else:
                return []
        except Exception as e:
            logging.error(f"Failed to extract work experience: {e}")
            return []
    
    def _extract_education(self, resume_text: str) -> List[Dict[str, Any]]:
        """Extract education from resume"""
        system_prompt = "You are a resume parser. Extract education information including dates."
        
        # Use full text for education extraction (education might be at the end)
        resume_snippet = resume_text
        
        user_prompt = f"""Extract the education from this resume. Look carefully for graduation dates/years which may appear as:
- Year ranges (e.g., 2018-2022)
- Single years (e.g., 2022)
- Month Year format (e.g., May 2022)
- "Expected" dates for ongoing education

Return ONLY valid JSON (use null for missing information):
{{
    "education": [
        {{
            "degree": "degree name",
            "institution": "school/university name",
            "graduation_date": "graduation year/date or date range if shown",
            "gpa": "GPA if mentioned or null",
            "honors": "honors/awards if mentioned or null"
        }}
    ]
}}

Resume text:
{resume_snippet}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            if not response or response == "{}":
                logging.warning("Empty response for education extraction")
                return []
            
            result = self._parse_llm_response(response)
            
            # Handle if response is dict with 'education' key
            if isinstance(result, dict):
                if 'education' in result:
                    education = result['education']
                    if education is None:
                        logging.warning("Education field is None in response")
                        return []
                    if isinstance(education, list):
                        return education
                    else:
                        logging.warning(f"Education field is not a list: {type(education)}")
                        return []
                else:
                    logging.warning("No 'education' key in response")
                    return []
            elif isinstance(result, list):
                return result
            else:
                logging.warning(f"Unexpected education response type: {type(result)}")
                return []
        except Exception as e:
            logging.error(f"Failed to extract education: {e}")
            return []
    
    def _extract_skills(self, resume_text: str) -> Dict[str, Any]:
        """Extract skills and certifications from resume"""
        system_prompt = "You are a resume parser. Extract skills and technologies."
        
        # Use full resume text for skills extraction
        resume_snippet = resume_text
        
        user_prompt = f"""Extract all skills from this resume. Return ONLY valid JSON:
{{
    "technical_skills": ["skill1", "skill2"],
    "soft_skills": ["skill1", "skill2"],
    "languages": ["language1", "language2"],
    "tools": ["tool1", "tool2"],
    "certifications": ["cert1", "cert2"]
}}

Resume text:
{resume_snippet}"""
        
        try:
            response = self._call_llm(system_prompt, user_prompt)
            result = self._parse_llm_response(response)
            # Map the response to expected format
            if isinstance(result, dict):
                return {
                    "technical_skills": result.get("technical_skills", []),
                    "programming_languages": result.get("languages", []),  # Map languages to programming_languages
                    "frameworks_tools": result.get("tools", []),  # Map tools to frameworks_tools
                    "certifications": result.get("certifications", [])
                }
            return {
                "technical_skills": [],
                "programming_languages": [],
                "frameworks_tools": [],
                "certifications": []
            }
        except Exception as e:
            logging.error(f"Failed to extract skills: {e}")
            return {
                "technical_skills": [],
                "programming_languages": [],
                "frameworks_tools": [],
                "certifications": []
            }
    
    def _safe_float_conversion(self, value) -> Optional[float]:
        """Safely convert a value to float, handling various edge cases"""
        if value is None:
            return None
        
        if isinstance(value, (int, float)):
            # Check for NaN or infinity
            if isinstance(value, float) and (value != value or value == float('inf') or value == float('-inf')):
                return None
            return float(value)
        
        if isinstance(value, str):
            try:
                result = float(value)
                # Check for NaN or infinity
                if result != result or result == float('inf') or result == float('-inf'):
                    return None
                return result
            except (ValueError, TypeError):
                return None
        
        return None
    
    def _parse_llm_response(self, response: str) -> Dict[str, Any]:
        """Parse LLM response with error handling"""
        
        if not response:
            logging.warning("Empty LLM response")
            return {}
        
        original_response = response
        response = response.strip()
        
        # Remove markdown code blocks if present
        if "```json" in response:
            response = response.split("```json")[1].split("```")[0]
        elif "```" in response:
            parts = response.split("```")
            if len(parts) >= 2:
                response = parts[1]
        
        response = response.strip()
        
        # First attempt: direct parse
        try:
            return json.loads(response)
        except json.JSONDecodeError as e:
            logging.debug(f"Initial parse failed at position {e.pos}: {e.msg}")
        
        # Second attempt: extract JSON pattern
        json_match = re.search(r'\{[^{}]*(?:\{[^{}]*(?:\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}[^{}]*)*\}[^{}]*)*\}', response, re.DOTALL)
        if json_match:
            try:
                json_str = json_match.group(0)
                # Clean up common issues
                json_str = re.sub(r',\s*]', ']', json_str)  # Remove trailing commas in arrays
                json_str = re.sub(r',\s*}', '}', json_str)  # Remove trailing commas in objects
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                logging.debug(f"Pattern extraction failed: {e}")
        
        # Third attempt: fix incomplete JSON
        if response.startswith('{'):
            # Count braces
            open_braces = response.count('{')
            close_braces = response.count('}')
            open_brackets = response.count('[')
            close_brackets = response.count(']')
            
            # Add missing closing characters
            if open_braces > close_braces:
                response += '}' * (open_braces - close_braces)
            if open_brackets > close_brackets:
                response += ']' * (open_brackets - close_brackets)
            
            try:
                return json.loads(response)
            except json.JSONDecodeError:
                pass
        
        # Log the problematic response for debugging
        logging.warning(f"Failed to parse LLM response. First 200 chars: {original_response[:200]}")
        
        # Return minimal structure
        return {
            "contact_info": {},
            "professional_summary": "",
            "technical_skills": [],
            "work_experience": [],
            "education": []
        }
    
    def extract_resume_info(self, resume_text: str) -> ResumeExtraction:
        """Extract structured information from resume text using multiple focused prompts"""
        try:
            logging.info("Starting multi-prompt resume extraction...")
            
            # Truncate resume if too long for overall processing
            if len(resume_text) > 10000:
                logging.info(f"Resume text too long ({len(resume_text)} chars), truncating to 10000")
                resume_text = resume_text[:10000]
            
            # Extract each section with focused prompts
            contact_info = self._extract_contact_info(resume_text) or {}
            summary = self._extract_summary(resume_text) or ""
            work_experience = self._extract_work_experience(resume_text) or []
            education = self._extract_education(resume_text) or []
            skills_data = self._extract_skills(resume_text) or {}
            
            # Log extraction results
            logging.info(f"Extraction results - Education items: {len(education) if education else 0}")
            logging.info(f"Extraction results - Work experience items: {len(work_experience) if work_experience else 0}")
            logging.info(f"Extraction results - Skills found: {bool(skills_data)}")
            
            # Combine all extracted data
            data = {
                "contact_info": contact_info,
                "professional_summary": summary,
                "work_experience": work_experience,
                "education": education,
                "technical_skills": skills_data.get("technical_skills", []),
                "programming_languages": skills_data.get("programming_languages", []),
                "frameworks_tools": skills_data.get("frameworks_tools", []),
                "certifications": skills_data.get("certifications", [])
            }
            
            # Process certifications - handle both list and nested format
            certs = data.get("certifications", [])
            if isinstance(certs, list) and len(certs) > 0:
                if isinstance(certs[0], str):
                    # Simple list of strings
                    certs = [{"name": cert, "issuer": None} for cert in certs]
                # else assume it's already in dict format
            
            # Calculate years of experience if not provided
            years_exp = data.get("years_of_experience")
            if years_exp is None or years_exp == "calculate_from_work_history":
                work_exp = data.get("work_experience", [])
                if work_exp:
                    # Try to calculate from first and last job dates
                    try:
                        first_job = work_exp[-1] if work_exp else None
                        last_job = work_exp[0] if work_exp else None
                        if first_job and last_job:
                            # Simple estimation - can be improved
                            years_exp = len(work_exp) * 2.5  # Rough estimate
                    except:
                        years_exp = None
            
            # Ensure all fields have proper defaults
            extraction_data = {
                "contact_info": data.get("contact_info", {}),
                "professional_summary": str(data.get("professional_summary", ""))[:1000],
                "technical_skills": [str(s) for s in data.get("technical_skills", []) if s],
                "programming_languages": [str(s) for s in data.get("programming_languages", []) if s],
                "frameworks_tools": [str(s) for s in data.get("frameworks_tools", []) if s],
                "work_experience": data.get("work_experience") or [],
                "education": data.get("education") or [],
                "certifications": certs or [],
                "achievements": data.get("achievements") or [],
                "years_of_experience": self._safe_float_conversion(years_exp),
                "salary_expectations": data.get("salary_expectations"),
                "location_preferences": data.get("location_preferences"),
                "work_authorization": data.get("work_authorization"),
                "availability": data.get("availability")
            }
            
            logging.info(f"Extraction complete - found {len(extraction_data['work_experience'])} jobs, {len(extraction_data['technical_skills'])} skills")
            
            return ResumeExtraction(**extraction_data)
            
        except Exception as e:
            logging.error(f"Resume extraction failed: {e}")
            import traceback
            traceback.print_exc()
            # Return minimal extraction
            return ResumeExtraction()


# -------------------------
# Resume Processing Pipeline
# -------------------------

class ResumeProcessingPipeline:
    """Main pipeline for processing resumes"""
    
    def __init__(
        self,
        config: ResumeConfig = None,
        llm_processor: ResumeLLMProcessor = None
    ):
        self.config = config or ResumeConfig()
        self.llm_processor = llm_processor or ResumeLLMProcessor(config=self.config)
        self.text_extractor = ResumeTextExtractor()
    
    def process_resume_file(
        self,
        file_path: str,
        verbose: bool = False
    ) -> Tuple[ProcessedResume, Dict[str, Any]]:
        """Process a resume file"""
        
        diagnostics = {
            "file_path": file_path,
            "file_exists": Path(file_path).exists(),
            "errors": []
        }
        
        try:
            # Extract text
            if verbose:
                logging.info(f"Extracting text from {file_path}")
            resume_text = self.text_extractor.extract_from_file(file_path)
            diagnostics["text_length"] = len(resume_text)
            
            # Process with LLM
            if verbose:
                logging.info("Processing with LLM...")
            extraction = self.llm_processor.extract_resume_info(resume_text)
            
            # Convert to ProcessedResume
            processed = self._build_processed_resume(extraction, resume_text)
            
            # Add metadata
            processed.metadata = {
                "file_path": file_path,
                "processed_at": datetime.now().isoformat(),
                "text_length": len(resume_text)
            }
            
            diagnostics["status"] = "success"
            diagnostics["sections_extracted"] = len([
                s for s in [
                    processed.technical_skills,
                    processed.work_experience,
                    processed.education
                ] if s
            ])
            
            return processed, diagnostics
            
        except Exception as e:
            logging.error(f"Resume processing failed: {e}")
            diagnostics["status"] = "error"
            diagnostics["errors"].append(str(e))
            
            # Return empty processed resume
            return ProcessedResume(
                contact_info=ContactInfo(),
                professional_summary="",
                technical_skills=[],
                programming_languages=[],
                frameworks_tools=[],
                work_experience=[],
                education=[],
                certifications=[],
                achievements=[]
            ), diagnostics
    
    def process_resume_directory(
        self,
        directory_path: str,
        extensions: List[str] = ['.txt', '.pdf', '.docx'],
        verbose: bool = False
    ) -> Tuple[List[ProcessedResume], Dict[str, Any]]:
        """Process all resumes in a directory"""
        
        directory = Path(directory_path)
        if not directory.exists():
            raise ValueError(f"Directory not found: {directory_path}")
        
        resumes = []
        diagnostics = {
            "directory": directory_path,
            "files_processed": 0,
            "files_failed": 0,
            "errors": []
        }
        
        # Find all resume files
        resume_files = []
        for ext in extensions:
            resume_files.extend(directory.glob(f"*{ext}"))
        
        diagnostics["files_found"] = len(resume_files)
        
        for file_path in resume_files:
            if verbose:
                logging.info(f"Processing {file_path}")
            
            try:
                processed, file_diag = self.process_resume_file(str(file_path), verbose)
                resumes.append(processed)
                diagnostics["files_processed"] += 1
            except Exception as e:
                logging.error(f"Failed to process {file_path}: {e}")
                diagnostics["files_failed"] += 1
                diagnostics["errors"].append(f"{file_path}: {str(e)}")
        
        return resumes, diagnostics
    
    def _build_processed_resume(
        self,
        extraction: ResumeExtraction,
        raw_text: str
    ) -> ProcessedResume:
        """Convert extraction to ProcessedResume"""
        
        # Build contact info
        contact_info = ContactInfo(**extraction.contact_info)
        
        # Add fallback for missing name
        if not contact_info.name or contact_info.name == "":
            # Try to generate a name from work experience or field
            if extraction.work_experience and len(extraction.work_experience) > 0:
                first_job = extraction.work_experience[0]
                job_title = first_job.get("title", "")
                if job_title:
                    contact_info.name = f"{job_title} Candidate"
                else:
                    contact_info.name = "Experienced Professional"
            else:
                # Use professional summary or skills to infer field
                if "payroll" in extraction.professional_summary.lower():
                    contact_info.name = "Payroll Professional"
                elif extraction.technical_skills:
                    contact_info.name = f"{extraction.technical_skills[0]} Specialist"
                else:
                    contact_info.name = "Professional Candidate"
        
        # Build work experience
        work_experience = []
        for exp in extraction.work_experience:
            work_experience.append(WorkExperience(
                title=exp.get("title", ""),
                company=exp.get("company", ""),
                location=exp.get("location"),
                start_date=exp.get("start_date"),
                end_date=exp.get("end_date"),
                duration=exp.get("duration"),
                responsibilities=exp.get("responsibilities", []),
                achievements=exp.get("achievements", []),
                technologies=exp.get("technologies", [])
            ))
        
        # Build education
        education = []
        for edu in extraction.education:
            education.append(Education(
                degree=edu.get("degree", ""),
                institution=edu.get("institution", ""),
                location=edu.get("location"),
                graduation_date=edu.get("graduation_date"),
                gpa=edu.get("gpa"),
                relevant_courses=edu.get("relevant_courses", [])
            ))
        
        # Build certifications
        certifications = []
        for cert in extraction.certifications:
            certifications.append(Certification(
                name=cert.get("name", ""),
                issuer=cert.get("issuer"),
                date_obtained=cert.get("date_obtained"),
                expiry_date=cert.get("expiry_date"),
                credential_id=cert.get("credential_id")
            ))
        
        return ProcessedResume(
            contact_info=contact_info,
            professional_summary=extraction.professional_summary,
            technical_skills=extraction.technical_skills,
            programming_languages=extraction.programming_languages,
            frameworks_tools=extraction.frameworks_tools,
            work_experience=work_experience,
            education=education,
            certifications=certifications,
            achievements=extraction.achievements,
            years_of_experience=extraction.years_of_experience,
            salary_expectations=extraction.salary_expectations,
            location_preferences=extraction.location_preferences,
            work_authorization=extraction.work_authorization,
            availability=extraction.availability,
            raw_text=raw_text
        )


# -------------------------
# Job Matching Functions
# -------------------------

def convert_resume_to_vocation_template(
    processed_resume: ProcessedResume,
    preferences: Optional[JobSeekerPreferences] = None
) -> VocationTemplate:
    """Convert processed resume to universal template for bidirectional matching"""
    
    # Create base template
    template = VocationTemplate(
        source_type="resume",
        source_id=processed_resume.contact_info.email or f"resume_{datetime.now().timestamp()}",
        title=processed_resume.work_experience[0].title if processed_resume.work_experience else "Professional",
        summary=processed_resume.professional_summary
    )
    
    # Convert technical skills
    for skill in processed_resume.technical_skills:
        template.technical_skills.append(SkillRequirement(
            skill_name=skill,
            required_proficiency="expert" if processed_resume.years_of_experience and processed_resume.years_of_experience > 5 else "advanced",
            years_required=None,
            is_mandatory=False,  # For resumes, these are skills they have
            evidence=[processed_resume.professional_summary[:100]],
            confidence=0.9
        ))
    
    # Convert programming languages and tools
    for lang in processed_resume.programming_languages:
        template.tools_technologies.append(SkillRequirement(
            skill_name=lang,
            required_proficiency="advanced",
            is_mandatory=False,
            confidence=0.9
        ))
    
    for tool in processed_resume.frameworks_tools:
        template.tools_technologies.append(SkillRequirement(
            skill_name=tool,
            required_proficiency="intermediate",
            is_mandatory=False,
            confidence=0.8
        ))
    
    # Add certifications
    template.certifications = [cert.name for cert in processed_resume.certifications]
    
    # Set experience
    template.total_years_experience = processed_resume.years_of_experience
    
    # Extract experience types from work history
    for exp in processed_resume.work_experience:
        if "manager" in exp.title.lower() or "lead" in exp.title.lower():
            if not template.management_experience:
                template.management_experience = 2.0  # Estimate
        
        # Add as experience requirement (what they bring)
        template.experience_requirements.append(ExperienceRequirement(
            experience_type="professional",
            years_required=0,  # They have this experience
            description=f"{exp.title} at {exp.company}",
            is_mandatory=False,
            evidence=exp.responsibilities[:2] if exp.responsibilities else [],
            confidence=1.0  # Direct from resume
        ))
    
    # Convert education
    for edu in processed_resume.education:
        template.education_requirements.append(EducationRequirement(
            degree_level=edu.degree,
            field_of_study=edu.institution,  # Using institution as field for now
            is_mandatory=False,
            confidence=1.0
        ))
    
    # Set location preferences
    location_pref = processed_resume.location_preferences or {}
    work_auth = processed_resume.work_authorization or ""
    
    template.location = LocationRequirement(
        preferred_locations=[processed_resume.contact_info.location] if processed_resume.contact_info.location else [],
        work_arrangement=_parse_candidate_work_preference(location_pref, preferences),
        relocation_willing=location_pref.get("relocation_willing", False) if isinstance(location_pref, dict) else False,
        visa_sponsorship_needed="visa" in work_auth.lower() if work_auth else False,
        work_authorization=work_auth
    )
    
    # Set compensation expectations
    salary_exp = processed_resume.salary_expectations or {}
    if isinstance(salary_exp, dict):
        template.compensation = CompensationRequirement(
            minimum_salary=salary_exp.get("minimum"),
            maximum_salary=salary_exp.get("maximum"),
            currency=salary_exp.get("currency", "USD"),
            benefits_required=salary_exp.get("benefits", [])
        )
    elif preferences and preferences.minimum_acceptable_salary:
        template.compensation.minimum_salary = preferences.minimum_acceptable_salary
    
    # Set culture fit from preferences
    if preferences:
        template.culture_fit = CultureFit(
            company_values=[],
            work_life_balance="important" if preferences.work_life_balance_importance > 0.7 else "flexible",
            team_size_preference=preferences.preferred_team_size,
            career_growth_importance="high" if preferences.career_growth_importance > 0.7 else "moderate",
            industry_preferences=preferences.preferred_industries,
            company_size_preference=preferences.preferred_company_size
        )
    
    # Add achievements
    all_achievements = processed_resume.achievements[:]
    for exp in processed_resume.work_experience:
        all_achievements.extend(exp.achievements)
    template.achievements = all_achievements[:10]  # Top 10
    
    # Add key responsibilities (what they've done)
    all_responsibilities = []
    for exp in processed_resume.work_experience:
        all_responsibilities.extend(exp.responsibilities)
    template.key_responsibilities = all_responsibilities[:10]  # Top 10
    
    # Set metadata
    template.metadata = {
        "name": processed_resume.contact_info.name,
        "email": processed_resume.contact_info.email,
        "phone": processed_resume.contact_info.phone,
        "linkedin": processed_resume.contact_info.linkedin,
        "github": processed_resume.contact_info.github,
        "portfolio": processed_resume.contact_info.portfolio,
        "availability": processed_resume.availability,
        "original_metadata": processed_resume.metadata
    }
    
    template.extraction_confidence = 0.85  # High confidence for structured resume data
    template.processed_at = datetime.now().isoformat()
    
    return template


def _parse_candidate_work_preference(
    location_pref: Dict[str, Any],
    preferences: Optional[JobSeekerPreferences]
) -> Optional[WorkArrangement]:
    """Parse work arrangement preference from resume and preferences"""
    
    if preferences and preferences.preferred_work_arrangement:
        return preferences.preferred_work_arrangement
    
    if isinstance(location_pref, dict):
        pref_str = str(location_pref.get("work_arrangement", "")).lower()
        if "remote" in pref_str:
            return WorkArrangement.REMOTE
        elif "hybrid" in pref_str:
            return WorkArrangement.HYBRID
        elif "onsite" in pref_str or "office" in pref_str:
            return WorkArrangement.ONSITE
    
    return WorkArrangement.FLEXIBLE


def create_matchable_resume_template(processed_resume: ProcessedResume) -> Dict[str, Any]:
    """Legacy function - converts to universal template then to dict"""
    template = convert_resume_to_vocation_template(processed_resume)
    
    # Convert to dictionary format for backward compatibility
    return {
        "candidate_profile": {
            "name": template.metadata.get("name"),
            "email": template.metadata.get("email"),
            "location": template.location.preferred_locations[0] if template.location.preferred_locations else None,
            "years_experience": template.total_years_experience
        },
        "qualifications": {
            "skills": [s.skill_name for s in template.technical_skills + template.soft_skills],
            "experience_summary": template.summary,
            "key_responsibilities": template.key_responsibilities,
            "achievements": template.achievements
        },
        "requirements_match": {
            "technical_skills": [s.skill_name for s in template.technical_skills],
            "programming_languages": [s.skill_name for s in template.tools_technologies if "language" in s.skill_name.lower()],
            "frameworks_tools": [s.skill_name for s in template.tools_technologies],
            "years_experience": template.total_years_experience,
            "education": [
                {"degree": e.degree_level, "field": e.field_of_study}
                for e in template.education_requirements
            ],
            "certifications": template.certifications
        },
        "preferences": {
            "salary_expectations": {
                "minimum": template.compensation.minimum_salary,
                "maximum": template.compensation.maximum_salary
            },
            "location_preferences": template.location.preferred_locations,
            "work_authorization": template.location.work_authorization,
            "availability": template.metadata.get("availability")
        },
        "work_history": [
            {
                "title": exp.description.split(" at ")[0] if " at " in exp.description else exp.description,
                "company": exp.description.split(" at ")[1] if " at " in exp.description else "N/A",
                "duration": f"{exp.years_required} years" if exp.years_required else "Current",
                "key_technologies": []
            } for exp in template.experience_requirements[:5]
        ]
    }


# -------------------------
# LangGraph Node Interface
# -------------------------

def resume_processing_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    LangGraph-compatible node for resume processing.
    
    Expected state keys:
        - resume_directory (str): Path to directory containing resumes
        - resume_files (list, optional): Specific resume files to process
        - config (dict, optional): Configuration overrides
        - verbose (bool, optional): Enable verbose logging
    
    Returns state with:
        - processed_resumes: List of processed resume dictionaries
        - resume_templates: List of matchable templates
        - diagnostics: Processing diagnostics
    """
    
    # Set up logging
    if state.get("verbose", False):
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
    
    # Get configuration
    config = ResumeConfig()
    if "config" in state and isinstance(state["config"], dict):
        config = ResumeConfig.from_dict(state["config"])
    
    # Initialize processor
    llm_processor = None
    if "llm_processor" in state:
        llm_processor = state["llm_processor"]
    
    pipeline = ResumeProcessingPipeline(
        config=config,
        llm_processor=llm_processor
    )
    
    try:
        # Process resumes
        if "resume_files" in state:
            # Process specific files
            resumes = []
            diagnostics = {"files_processed": 0, "errors": []}
            
            for file_path in state["resume_files"]:
                try:
                    processed, file_diag = pipeline.process_resume_file(
                        file_path,
                        verbose=state.get("verbose", False)
                    )
                    resumes.append(processed)
                    diagnostics["files_processed"] += 1
                except Exception as e:
                    diagnostics["errors"].append(f"{file_path}: {str(e)}")
        else:
            # Process directory
            resume_dir = state.get("resume_directory", "data/resumes")
            resumes, diagnostics = pipeline.process_resume_directory(
                resume_dir,
                verbose=state.get("verbose", False)
            )
        
        # Store in database if available
        db = state.get("database")
        if db:
            from graph.database import graphDB
            if not isinstance(db, graphDB):
                db = graphDB()
                state["database"] = db
        
        # Convert to dictionaries and create templates
        processed_resumes = []
        resume_templates = []
        
        for resume in resumes:
            # Extract filename from metadata
            filename = resume.metadata.get('file_path', 'unknown.txt')
            if '/' in filename:
                filename = filename.split('/')[-1]
            elif '\\' in filename:
                filename = filename.split('\\')[-1]
            
            # Convert to dictionary with all extracted data for database storage
            resume_dict = {
                "contact_info": {
                    "name": resume.contact_info.name,
                    "email": resume.contact_info.email,
                    "phone": resume.contact_info.phone,
                    "location": resume.contact_info.location,
                    "linkedin": resume.contact_info.linkedin,
                    "github": resume.contact_info.github,
                    "portfolio": resume.contact_info.portfolio
                },
                "professional_summary": resume.professional_summary,
                "technical_skills": resume.technical_skills,
                "programming_languages": resume.programming_languages,
                "frameworks_tools": resume.frameworks_tools,
                "work_experience": [
                    {
                        "title": exp.title,
                        "company": exp.company,
                        "location": exp.location,
                        "start_date": exp.start_date,
                        "end_date": exp.end_date,
                        "duration": exp.duration,
                        "responsibilities": exp.responsibilities,
                        "achievements": exp.achievements,
                        "technologies": exp.technologies
                    } for exp in resume.work_experience
                ],
                "education": [
                    {
                        "degree": edu.degree,
                        "institution": edu.institution,
                        "location": edu.location,
                        "graduation_date": edu.graduation_date,
                        "gpa": edu.gpa,
                        "relevant_courses": edu.relevant_courses
                    } for edu in resume.education
                ],
                "certifications": [
                    {
                        "name": cert.name,
                        "issuer": cert.issuer,
                        "date_obtained": cert.date_obtained,
                        "expiry_date": cert.expiry_date,
                        "credential_id": cert.credential_id
                    } for cert in resume.certifications
                ],
                "achievements": resume.achievements,
                "years_of_experience": resume.years_of_experience,
                "salary_expectations": resume.salary_expectations,
                "location_preferences": resume.location_preferences,
                "work_authorization": resume.work_authorization,
                "availability": resume.availability,
                "full_text": resume.raw_text,  # Include full text for embeddings
                "filename": filename,
                "metadata": resume.metadata
            }
            processed_resumes.append(resume_dict)
            
            # Create matchable template
            template = create_matchable_resume_template(resume)
            resume_dict["matching_template"] = template
            
            # Create universal template
            vocation_template = convert_resume_to_vocation_template(resume)
            # Convert to dict, handling enums properly
            universal_dict = asdict(vocation_template)
            # Convert any enum values to strings
            def convert_enums(obj):
                if isinstance(obj, dict):
                    return {k: convert_enums(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [convert_enums(item) for item in obj]
                elif hasattr(obj, 'value'):  # Enum
                    return obj.value
                else:
                    return obj
            resume_dict["vocation_template"] = convert_enums(universal_dict)
            resume_templates.append(template)
            
            # Store in database with section embeddings
            if db:
                try:
                    resume_id = db.add_resume(resume_dict)
                    resume_dict["resume_id"] = resume_id
                    logger.info(f"Stored resume {filename} in database with ID: {resume_id}")
                except Exception as e:
                    logger.error(f"Failed to store resume in database: {e}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
        
        # Update state
        state["processed_resumes"] = processed_resumes
        state["resume_templates"] = resume_templates
        
        if "diagnostics" not in state:
            state["diagnostics"] = {}
        state["diagnostics"]["resume_processing"] = diagnostics
        
    except Exception as e:
        logging.error(f"Resume processing failed: {e}")
        state["processed_resumes"] = []
        state["resume_templates"] = []
        state["diagnostics"] = {
            "resume_processing": {
                "status": "error",
                "error": str(e)
            }
        }
    
    return state


# -------------------------
# Standalone Testing
# -------------------------

if __name__ == "__main__":
    # Test with sample resume
    import sys
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    
    # Check if resume file provided
    if len(sys.argv) > 1:
        resume_file = sys.argv[1]
    else:
        # Use default test file
        resume_file = "data/resumes/naloni_resume_20250812.txt"
    
    print(f"Processing resume: {resume_file}")
    print("-" * 60)
    
    # Test pipeline
    pipeline = ResumeProcessingPipeline()
    
    try:
        processed, diagnostics = pipeline.process_resume_file(resume_file, verbose=True)
        
        print("\n" + "=" * 60)
        print("EXTRACTED INFORMATION")
        print("=" * 60)
        
        # Contact Info
        print(f"\nContact Information:")
        print(f"  Name: {processed.contact_info.name}")
        print(f"  Email: {processed.contact_info.email}")
        print(f"  Location: {processed.contact_info.location}")
        
        # Summary
        print(f"\nProfessional Summary:")
        print(f"  {processed.professional_summary[:200]}...")
        
        # Skills
        print(f"\nTechnical Skills ({len(processed.technical_skills)}):")
        for skill in processed.technical_skills[:10]:
            print(f"  - {skill}")
        
        # Experience
        print(f"\nWork Experience ({len(processed.work_experience)} positions):")
        for exp in processed.work_experience[:3]:
            print(f"\n  {exp.title} at {exp.company}")
            print(f"  {exp.start_date} - {exp.end_date}")
            if exp.responsibilities:
                print(f"  Key responsibilities:")
                for resp in exp.responsibilities[:3]:
                    print(f"    - {resp}")
        
        # Education
        print(f"\nEducation ({len(processed.education)}):")
        for edu in processed.education:
            print(f"  {edu.degree} from {edu.institution}")
        
        # Create matchable template
        print("\n" + "=" * 60)
        print("MATCHABLE TEMPLATE")
        print("=" * 60)
        
        template = create_matchable_resume_template(processed)
        print(json.dumps(template, indent=2)[:1000] + "...")
        
        print(f"\nDiagnostics:")
        print(json.dumps(diagnostics, indent=2))
        
    except Exception as e:
        print(f"Error processing resume: {e}")
        import traceback
        traceback.print_exc()
